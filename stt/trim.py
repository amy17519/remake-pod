"""
Podcast Transcript Trimmer

Uses Claude to identify repetitive, non-insightful, or setup content from podcast
transcripts. Supports two types of edits:
  - Whole-segment cuts: entire time ranges removed (pre-roll, filler, off-topic)
  - Intra-segment edits: specific phrases within a segment wrapped in [DEL]...[/DEL]
    to trim rambling guest answers while preserving the key point

Outputs a .docx with deletions shown in red strikethrough, plus a cuts JSON.

Prerequisites:
    ANTHROPIC_API_KEY environment variable must be set
    pip install anthropic python-docx

Usage:
    python trim.py <transcript.txt> [--audio AUDIO_FILE] [--output-dir DIR]

Arguments:
    transcript          Path to transcript .txt file (from rev.py output)
    --audio, -a         Path to audio file to generate ffmpeg cut command
    --output-dir, -o    Directory for output files (default: ./results)

Output:
    results/{name}_trimmed.docx - Transcript with deletions in red strikethrough
    results/{name}_cuts.json    - Cut decisions with timestamps and reasons
    If --audio is provided: prints ffmpeg command to cut the audio file
"""

import anthropic
import argparse
import json
import os
import re

from docx import Document
from docx.shared import RGBColor


def parse_transcript(text):
    """Parse transcript into list of segment dicts: {speaker, timestamp, content}."""
    segments = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if i + 1 < len(lines) and re.match(r"^\d{2}:\d{2}:\d{2}$", lines[i + 1].strip()):
            speaker = line
            timestamp = lines[i + 1].strip()
            content = lines[i + 2].strip() if i + 2 < len(lines) else ""
            segments.append({"speaker": speaker, "timestamp": timestamp, "content": content})
            i += 3
        else:
            i += 1
    return segments


def timestamp_to_seconds(ts):
    h, m, s = ts.split(":")
    return int(h) * 3600 + int(m) * 60 + int(s)


def seconds_to_hms(total):
    h = total // 3600
    m = (total % 3600) // 60
    s = total % 60
    return f"{h:02d}:{m:02d}:{s:02d}"


def format_for_claude(segments):
    return "\n".join(f"[{s['timestamp']}] {s['speaker']}: {s['content']}" for s in segments)


HOST_NAMES = {"amy", "stella"}


def split_into_chunks(segments, target_minutes=30, max_overshoot_minutes=5):
    """
    Split segments into ~target_minutes chunks, breaking at natural topic boundaries.
    After reaching target_minutes, waits for a point where the next segment is a host
    question (Amy/Stella asking something) — so we never break mid-topic or mid-answer.
    Falls back to any speaker transition if no host question found within max_overshoot_minutes.
    """
    target_seconds = target_minutes * 60
    max_overshoot = max_overshoot_minutes * 60
    chunks = []
    i = 0

    while i < len(segments):
        chunk_start_t = timestamp_to_seconds(segments[i]["timestamp"])
        chunk = []
        in_overshoot = False

        while i < len(segments):
            seg = segments[i]
            t = timestamp_to_seconds(seg["timestamp"])
            elapsed = t - chunk_start_t

            if elapsed >= target_seconds:
                in_overshoot = True

            if in_overshoot and i + 1 < len(segments):
                next_seg = segments[i + 1]
                next_speaker = next_seg["speaker"].strip().lower()
                is_host_question = (
                    any(h in next_speaker for h in HOST_NAMES)
                    and "?" in next_seg["content"]
                )
                past_max = elapsed >= target_seconds + max_overshoot

                if is_host_question or past_max:
                    chunk.append(seg)
                    i += 1
                    break

            chunk.append(seg)
            i += 1

        if chunk:
            chunks.append(chunk)

    return chunks


def analyze_chunk(chunk_segments, client, chunk_num, total_chunks):
    """Analyze a single chunk and return cut/edit decisions."""
    start_ts = chunk_segments[0]["timestamp"]
    end_ts = chunk_segments[-1]["timestamp"]
    transcript_text = format_for_claude(chunk_segments)

    prompt = f"""You are a professional podcast editor working on chunk {chunk_num} of {total_chunks} ({start_ts} to {end_ts}). Cut this segment to roughly 50% of its length. You have two tools:"""

    return _call_claude(prompt, transcript_text, client)


def analyze_with_claude(transcript_text, client):
    """Send a full transcript to Claude and get cut decisions as JSON (single-pass mode)."""
    prompt = """You are a professional podcast editor. The raw episode is approximately 3 hours long. Cut it down to 1–1.5 hours. You have two tools:"""
    return _call_claude(prompt, transcript_text, client)


def _call_claude(prompt_header, transcript_text, client):
    """Shared Claude call used by both single-pass and chunked modes."""
    prompt = f"""{prompt_header}

1. WHOLE-SEGMENT CUTS ("cuts"): Remove an entire time range. Use for pre-recording setup, post-recording chatter, filler-only exchanges, entire off-topic tangents, and sections where the same point repeats across multiple speakers over several minutes.

2. INTRA-SEGMENT EDITS ("edits"): Edit inside a single guest segment by wrapping the parts to remove in [DEL]...[/DEL] tags. Use this when a guest gives a long answer that starts strong but then rambles, repeats, or loses the thread. Keep the sharpest version of their point; cut the padding. Make the guest sound concise and smart.

1. WHOLE-SEGMENT CUTS ("cuts"): Remove an entire time range. Use for pre-recording setup, post-recording chatter, filler-only exchanges, entire off-topic tangents, and sections where the same point repeats across multiple speakers over several minutes.

2. INTRA-SEGMENT EDITS ("edits"): Edit inside a single guest segment by wrapping the parts to remove in [DEL]...[/DEL] tags. Use this when a guest gives a long answer that starts strong but then rambles, repeats, or loses the thread. Keep the sharpest version of their point; cut the padding. Make the guest sound concise and smart.

NEVER cut:
- The podcast intro: Amy saying "哈喽大家好我是Amy" or "欢迎来到" — always keep
- Host questions (Amy or Stella asking something) — always keep, even short ones. Listener needs them to understand the guest's answers
- Personal/philosophical tangents about motherhood, raising kids in the AI era, happiness, society, nostalgia, life impact of AI — these are the heart of the podcast, keep them all
- Substantive guest insights, career stories, and personal opinions

WHOLE-SEGMENT CUT rules — cut without hesitation:
1. Pre-recording setup: QuickTime/Zoom setup, audio checks, clap syncs, mic/gear chat, topic negotiation — cut all
2. Post-recording chatter after sign-off
3. Segments that are purely 嗯 / 对对对 / 好的好的好的 / 嗯嗯嗯 with zero other content
4. Hosts looking things up live (searching episode numbers, pulling up links)
5. Off-topic host side chat irrelevant to listeners (mutual acquaintance name-dropping, inside jokes)
6. Stuttered false starts that get restated cleanly right after
7. Meta-commentary about podcast production (editing plans, question allocation, skipping sections)

INTRA-SEGMENT EDIT rules — apply to guest answers that:
- Repeat the same point 2–3 times within one answer
- Spend multiple sentences building up to a simple point that could be stated in one
- Trail off into hedging or self-correction before landing on the real answer
- Include filler phrases like "就是说", "我的意思是", "怎么说呢", "嗯 然后" mid-answer that break the flow

When editing intra-segment, use [DEL]...[/DEL] tags around the exact text to remove. Only edit the content field. Keep speaker and timestamp intact. The result should read naturally after the deletions — don't leave broken sentences.

Return JSON only, exactly this structure:
{{
  "cuts": [
    {{
      "start_time": "HH:MM:SS",
      "end_time": "HH:MM:SS",
      "reason": "brief reason"
    }}
  ],
  "edits": [
    {{
      "timestamp": "HH:MM:SS",
      "speaker": "speaker name",
      "content": "the full original content with [DEL]text to remove[/DEL] inline",
      "reason": "brief reason"
    }}
  ],
  "summary": "1-2 sentence summary"
}}

Transcript:
{transcript_text}"""

    full_response = ""
    with client.messages.stream(
        model="claude-opus-4-6",
        max_tokens=16384,
        messages=[{"role": "user", "content": prompt}],
    ) as stream:
        for text in stream.text_stream:
            full_response += text
            print(text, end="", flush=True)
    print()

    response = full_response.strip()
    response = re.sub(r"^```(?:json)?\n?", "", response)
    response = re.sub(r"\n?```$", "", response)

    return json.loads(response)


def process_in_chunks(segments, client, target_minutes=30):
    """Split transcript into ~30min chunks and process each independently."""
    chunks = split_into_chunks(segments, target_minutes=target_minutes)
    print(f"Split into {len(chunks)} chunks:")
    for i, chunk in enumerate(chunks):
        print(f"  Chunk {i+1}: {chunk[0]['timestamp']} → {chunk[-1]['timestamp']} ({len(chunk)} segments)")

    all_cuts = []
    all_edits = []
    summaries = []

    for i, chunk in enumerate(chunks):
        print(f"\n--- Chunk {i+1}/{len(chunks)} ({chunk[0]['timestamp']} → {chunk[-1]['timestamp']}) ---\n")
        result = analyze_chunk(chunk, client, i + 1, len(chunks))
        all_cuts.extend(result.get("cuts", []))
        all_edits.extend(result.get("edits", []))
        if result.get("summary"):
            summaries.append(f"[{chunk[0]['timestamp']}–{chunk[-1]['timestamp']}] {result['summary']}")
        n_cuts = len(result.get("cuts", []))
        n_edits = len(result.get("edits", []))
        print(f"\n  → {n_cuts} cuts, {n_edits} edits")

    return {
        "cuts": all_cuts,
        "edits": all_edits,
        "summary": " | ".join(summaries),
    }


def apply_cuts(segments, cuts_data):
    """
    Apply both whole-segment cuts and intra-segment edits.
    - deleted=True: entire segment is cut
    - edited_content: content string with [DEL]...[/DEL] markers for partial cuts
    """
    cut_ranges = [
        (timestamp_to_seconds(c["start_time"]), timestamp_to_seconds(c["end_time"]))
        for c in cuts_data.get("cuts", [])
    ]

    # Build a lookup: timestamp -> edited content
    edits_by_ts = {}
    for edit in cuts_data.get("edits", []):
        edits_by_ts[edit["timestamp"]] = edit["content"]

    result = []
    for seg in segments:
        seg = dict(seg)
        t = timestamp_to_seconds(seg["timestamp"])
        seg["deleted"] = any(start <= t <= end for start, end in cut_ranges)
        if not seg["deleted"] and seg["timestamp"] in edits_by_ts:
            seg["edited_content"] = edits_by_ts[seg["timestamp"]]
        result.append(seg)
    return result


def write_paragraph_with_dels(doc, text, is_deleted_segment=False):
    """
    Add a paragraph to doc. If is_deleted_segment, entire text is red strikethrough.
    Otherwise parse [DEL]...[/DEL] tags and apply red strikethrough only to those spans.
    """
    p = doc.add_paragraph()
    if is_deleted_segment:
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        run.font.strike = True
        return

    # Parse [DEL]...[/DEL] tags
    parts = re.split(r"(\[DEL\].*?\[/DEL\])", text, flags=re.DOTALL)
    for part in parts:
        if part.startswith("[DEL]") and part.endswith("[/DEL]"):
            inner = part[5:-6]
            run = p.add_run(inner)
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.font.strike = True
        else:
            if part:
                p.add_run(part)


def write_transcript_docx(segments, path):
    """Write transcript as .docx with deletions in red strikethrough."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = section.bottom_margin = 914400 // 2
    section.left_margin = section.right_margin = 914400

    for seg in segments:
        deleted = seg.get("deleted", False)
        edited_content = seg.get("edited_content")

        # Speaker line
        p = doc.add_paragraph()
        run = p.add_run(seg["speaker"])
        if deleted:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.font.strike = True

        # Timestamp line
        p = doc.add_paragraph()
        run = p.add_run(seg["timestamp"])
        if deleted:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.font.strike = True

        # Content line
        if deleted:
            write_paragraph_with_dels(doc, seg["content"], is_deleted_segment=True)
        elif edited_content:
            write_paragraph_with_dels(doc, edited_content, is_deleted_segment=False)
        else:
            doc.add_paragraph(seg["content"])

        doc.add_paragraph()  # blank line between segments

    doc.save(path)


def build_ffmpeg_command(audio_file, cuts_data, total_duration_s, output_path):
    cuts = sorted(cuts_data.get("cuts", []), key=lambda c: timestamp_to_seconds(c["start_time"]))

    keep_intervals = []
    cursor = 0
    for cut in cuts:
        cut_start = timestamp_to_seconds(cut["start_time"])
        cut_end = timestamp_to_seconds(cut["end_time"])
        if cursor < cut_start:
            keep_intervals.append((cursor, cut_start))
        cursor = max(cursor, cut_end)
    keep_intervals.append((cursor, None))

    if not keep_intervals:
        return None, []

    n = len(keep_intervals)
    filter_parts = []
    labels = []
    for i, (start, end) in enumerate(keep_intervals):
        trim = f"atrim=start={start}" + (f":end={end}" if end is not None else "")
        filter_parts.append(f"[0:a]{trim},asetpts=PTS-STARTPTS[a{i}]")
        labels.append(f"[a{i}]")
    filter_parts.append(f"{''.join(labels)}concat=n={n}:v=0:a=1[out]")

    filter_complex = "; ".join(filter_parts)
    cmd = f'ffmpeg -i "{audio_file}" -filter_complex "{filter_complex}" -map "[out]" "{output_path}"'
    return cmd, keep_intervals


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Trim podcast transcript using Claude")
    parser.add_argument("transcript", help="Path to transcript .txt file")
    parser.add_argument("--audio", "-a", help="Path to audio file (generates ffmpeg command)")
    parser.add_argument("--output-dir", "-o", default="./results", help="Output directory (default: ./results)")
    parser.add_argument("--chunk-minutes", "-c", type=int, default=30, help="Chunk size in minutes (default: 30). Use 0 to process all at once.")
    args = parser.parse_args()

    if not os.path.exists(args.transcript):
        print(f"Error: {args.transcript} not found")
        exit(1)

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        print("Error: ANTHROPIC_API_KEY environment variable not set")
        exit(1)

    client = anthropic.Anthropic(api_key=api_key)

    print(f"Reading transcript: {args.transcript}")
    with open(args.transcript, "r", encoding="utf-8") as f:
        text = f.read()

    segments = parse_transcript(text)
    print(f"Parsed {len(segments)} segments\n")

    if args.chunk_minutes == 0:
        print("Processing full transcript in one pass...\n")
        cuts_data = analyze_with_claude(format_for_claude(segments), client)
    else:
        print(f"Processing in ~{args.chunk_minutes}-minute chunks...\n")
        cuts_data = process_in_chunks(segments, client, target_minutes=args.chunk_minutes)

    n_cuts = len(cuts_data.get("cuts", []))
    n_edits = len(cuts_data.get("edits", []))
    print(f"\nWhole-segment cuts: {n_cuts}, Intra-segment edits: {n_edits}")
    if cuts_data.get("summary"):
        print(f"Summary: {cuts_data['summary']}\n")

    base = os.path.splitext(os.path.basename(args.transcript))[0]
    os.makedirs(args.output_dir, exist_ok=True)

    cuts_path = os.path.join(args.output_dir, f"{base}_cuts.json")
    with open(cuts_path, "w", encoding="utf-8") as f:
        json.dump(cuts_data, f, ensure_ascii=False, indent=2)
    print(f"Cut decisions saved to: {cuts_path}")

    trimmed = apply_cuts(segments, cuts_data)
    trimmed_path = os.path.join(args.output_dir, f"{base}_trimmed.docx")
    write_transcript_docx(trimmed, trimmed_path)

    n_deleted = sum(1 for s in trimmed if s.get("deleted"))
    n_edited = sum(1 for s in trimmed if s.get("edited_content"))
    print(f"Trimmed transcript saved to: {trimmed_path}")
    print(f"Segments: {len(segments)} total | {n_deleted} fully cut ({n_deleted*100//len(segments)}%) | {n_edited} partially edited")

    if args.audio:
        if not os.path.exists(args.audio):
            print(f"\nWarning: audio file not found at {args.audio}")
        else:
            ext = os.path.splitext(args.audio)[1]
            audio_out = os.path.join(args.output_dir, f"{base}_trimmed{ext}")
            cmd, keep_intervals = build_ffmpeg_command(args.audio, cuts_data, None, audio_out)
            if cmd:
                print(f"\nffmpeg command to cut audio:")
                print(f"  {cmd}")
                print(f"\nKeeping {len(keep_intervals)} interval(s):")
                for start, end in keep_intervals:
                    end_str = seconds_to_hms(end) if end is not None else "end"
                    print(f"  {seconds_to_hms(start)} → {end_str}")
